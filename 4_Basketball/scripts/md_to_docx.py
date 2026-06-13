"""Convert a markdown file to .docx.

Prefers pandoc (via pypandoc) for high-fidelity conversion (tables, code
blocks, edge cases). Falls back to a python-docx implementation if pandoc
is unavailable.

Usage:
  python scripts/md_to_docx.py path/to/file.md
  python scripts/md_to_docx.py path/to/file.md --out path/to/output.docx
  python scripts/md_to_docx.py path/to/file.md --backend python-docx

Default output is alongside the input with .docx extension.
"""
from __future__ import annotations
import argparse
import re
import sys
from pathlib import Path

try:
    import pypandoc
    HAS_PANDOC = True
except Exception:
    HAS_PANDOC = False

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def md_to_docx_pandoc(md_path: Path, out_path: Path):
    """High-fidelity conversion via pandoc."""
    pypandoc.convert_file(
        str(md_path), "docx", outputfile=str(out_path),
        extra_args=["--standalone"],
    )
    return out_path


INLINE_RE = re.compile(
    r"(\*\*([^*]+)\*\*|"             # bold
    r"`([^`]+)`|"                    # inline code
    r"\*([^*]+)\*|"                   # italic
    r"\[([^\]]+)\]\(([^)]+)\))"      # link
)


def add_inline(paragraph, text):
    """Add a string with inline markdown to a paragraph as runs."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        pre = text[pos:m.start()]
        if pre:
            paragraph.add_run(pre)
        if m.group(2) is not None:
            r = paragraph.add_run(m.group(2))
            r.bold = True
        elif m.group(3) is not None:
            r = paragraph.add_run(m.group(3))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif m.group(4) is not None:
            r = paragraph.add_run(m.group(4))
            r.italic = True
        elif m.group(5) is not None:
            r = paragraph.add_run(m.group(5))
            r.italic = True
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        pos = m.end()
    rest = text[pos:]
    if rest:
        paragraph.add_run(rest)


def parse_table(lines, start_idx):
    """Parse a markdown pipe-table starting at start_idx.
    Returns (table_data, end_idx) where end_idx is one past the last table row.
    """
    rows = []
    i = start_idx
    while i < len(lines) and "|" in lines[i]:
        line = lines[i].strip()
        if not line:
            break
        # Skip the separator row
        if re.match(r"^\|?[\s:|-]+\|?$", line) and "-" in line:
            i += 1
            continue
        # Strip leading/trailing pipes; split
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def md_to_docx(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # Set base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            if in_code_block:
                # Close code block
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buffer))
                r.font.name = "Consolas"
                r.font.size = Pt(9)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.add_run("─" * 60)
            i += 1
            continue

        # Headings
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            txt = line.lstrip("#").strip()
            if level <= 4:
                doc.add_heading(txt, level=level)
            else:
                p = doc.add_paragraph()
                r = p.add_run(txt)
                r.bold = True
            i += 1
            continue

        # Pipe table
        if line.strip().startswith("|") or (i + 1 < len(lines)
                                              and "|" in line
                                              and re.match(r"^\|?[\s:|-]+\|?$",
                                                            lines[i + 1].strip())
                                              and "-" in lines[i + 1]):
            rows, new_i = parse_table(lines, i)
            if rows:
                ncols = max(len(r) for r in rows)
                t = doc.add_table(rows=len(rows), cols=ncols)
                t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci in range(ncols):
                        cell = t.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        cell_text = row[ci] if ci < len(row) else ""
                        add_inline(p, cell_text)
                        if ri == 0:
                            for run in p.runs:
                                run.bold = True
                i = new_i
                doc.add_paragraph()
                continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            txt = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, txt)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            txt = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, txt)
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            txt = line.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r_marker = p.add_run("│ ")
            r_marker.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            add_inline(p, txt)
            i += 1
            continue

        # Blank line → blank paragraph
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Plain paragraph (collect consecutive non-blank, non-special lines)
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if not nxt.strip(): break
            if nxt.startswith("#"): break
            if nxt.strip().startswith("```"): break
            if nxt.strip().startswith("|"): break
            if re.match(r"^\s*[-*]\s+", nxt): break
            if re.match(r"^\s*\d+\.\s+", nxt): break
            if nxt.startswith(">"): break
            if nxt.strip() in ("---", "***", "___"): break
            para_lines.append(nxt)
            j += 1
        para_text = " ".join(p.strip() for p in para_lines)
        p = doc.add_paragraph()
        add_inline(p, para_text)
        i = j

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md_path", type=str)
    ap.add_argument("--out", type=str, default=None)
    ap.add_argument("--backend", choices=["pandoc", "python-docx", "auto"],
                     default="auto")
    args = ap.parse_args()

    md_path = Path(args.md_path)
    out_path = Path(args.out) if args.out else md_path.with_suffix(".docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    backend = args.backend
    if backend == "auto":
        backend = "pandoc" if HAS_PANDOC else "python-docx"

    print(f"[convert via {backend}] {md_path}  ->  {out_path}")
    if backend == "pandoc":
        md_to_docx_pandoc(md_path, out_path)
    else:
        md_to_docx(md_path, out_path)
    print(f"[done] {out_path}")


if __name__ == "__main__":
    main()
